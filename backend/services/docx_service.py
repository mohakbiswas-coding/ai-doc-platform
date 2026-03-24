# backend/services/docx_service.py

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


def build_docx(project_title: str, topic: str, sections: list) -> bytes:
    """
    Build a .docx file from project sections.
    Returns file as bytes.
    """
    doc = Document()

    # ── Title page ────────────────────────────────────────────────────────────
    title_para = doc.add_heading(project_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    topic_para = doc.add_paragraph(f"Topic: {topic}")
    topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── Table of Contents (simple text list) ──────────────────────────────────
    doc.add_heading("Table of Contents", level=1)
    for section in sections:
        doc.add_paragraph(section["title"], style="List Number")
    doc.add_page_break()

    # ── Sections ──────────────────────────────────────────────────────────────
    for section in sections:
        # Section heading
        doc.add_heading(section["title"], level=1)

        # Section content
        content = section.get("content", "")
        if content:
            for paragraph in content.split("\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    doc.add_paragraph(paragraph)

        doc.add_paragraph()  # Blank line between sections

    # ── Save to bytes ─────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
